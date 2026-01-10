import sys
import os
import json
import shutil
import subprocess
import markdown
import pypandoc
import re
import ctypes
from datetime import datetime
import tempfile
import xml.etree.ElementTree as etree

# --- Markdown 扩展依赖 ---
from markdown.inlinepatterns import InlineProcessor
from markdown.extensions import Extension

from PyQt6.QtWidgets import (QApplication, QMainWindow, QWidget, QVBoxLayout,
                             QHBoxLayout, QTextEdit, QFileDialog, QPushButton,
                             QMessageBox, QSplitter, QMenu, QFrame, QDialog,
                             QLabel, QDoubleSpinBox, QCheckBox, QGroupBox,
                             QLineEdit, QSpinBox, QFormLayout, QDialogButtonBox)
from PyQt6.QtWebEngineCore import QWebEnginePage, QWebEngineSettings
from PyQt6.QtWebEngineWidgets import QWebEngineView
from PyQt6.QtGui import QAction, QFont, QIcon, QTextCursor, QPageLayout, QPageSize, QKeySequence, QShortcut
from PyQt6.QtCore import QUrl, Qt, QTimer, QPoint, QMarginsF, QMimeData

# =========================================================================
# [全局配置]
# =========================================================================

if getattr(sys, 'frozen', False):
    BASE_DIR = os.path.dirname(sys.executable)
    if hasattr(sys, '_MEIPASS'):
        TEMP_RESOURCE_DIR = sys._MEIPASS
    else:
        TEMP_RESOURCE_DIR = BASE_DIR
else:
    BASE_DIR = os.path.dirname(os.path.abspath(__file__))
    TEMP_RESOURCE_DIR = BASE_DIR

SETTINGS_FILE = os.path.join(BASE_DIR, "settings.json")


def resource_path(relative_path):
    return os.path.join(TEMP_RESOURCE_DIR, relative_path)


# =========================================================================
# [调试] 控制台日志重定向
# =========================================================================
class DebugWebPage(QWebEnginePage):
    def javaScriptConsoleMessage(self, level, msg, line, sourceID):
        if "Error" in msg or "Fail" in msg:
            print(f"[JS Error] {msg} (Line {line})")


# =========================================================================
# [核心逻辑] 强力公式保护扩展
# =========================================================================
class MathProcessor(InlineProcessor):
    def __init__(self, pattern, md, mode):
        super().__init__(pattern, md)
        self.mode = mode

    def handleMatch(self, m, data):
        tex_content = m.group(1)
        if self.mode == 'block':
            placeholder = self.md.htmlStash.store(
                f'<div class="katex-raw" data-display="true">{tex_content}</div>'
            )
        else:
            placeholder = self.md.htmlStash.store(
                f'<span class="katex-raw" data-display="false">{tex_content}</span>'
            )
        return placeholder, m.start(0), m.end(0)


class MathExtension(Extension):
    def extendMarkdown(self, md):
        md.inlinePatterns.register(
            MathProcessor(r'(?<!\\)\$\$([\s\S]+?)(?<!\\)\$\$', md, 'block'),
            'math_block', 185
        )
        md.inlinePatterns.register(
            MathProcessor(r'(?<!\\)\$([^$]+?)(?<!\\)\$', md, 'inline'),
            'math_inline', 185
        )


# =========================================================================
# [核心逻辑]
# =========================================================================
def load_settings_raw():
    if os.path.exists(SETTINGS_FILE):
        try:
            with open(SETTINGS_FILE, 'r', encoding='utf-8') as f:
                return json.load(f)
        except:
            pass
    return {}


def save_settings_raw(data):
    try:
        current = load_settings_raw()
        current.update(data)
        with open(SETTINGS_FILE, 'w', encoding='utf-8') as f:
            json.dump(current, f, indent=4)
    except Exception as e:
        print(f"保存配置失败: {e}")


# =========================================================================
# Pandoc 安装检测与自动安装
# =========================================================================
def check_pandoc_dependency():
    # 1. 优先检查系统路径中是否真的存在 pandoc
    # 无论 settings.json 怎么写，实测为准
    is_system_installed = shutil.which("pandoc") is not None

    if is_system_installed:
        # A. 如果系统已安装，确保配置文件标记为 True
        save_settings_raw({"pandoc_installed": True})

        # B. [新增功能] 自动删除安装包以释放空间
        try:
            msi_path = resource_path("pandoc.msi")
            # 只有在打包模式下，且文件存在时才删除
            if getattr(sys, 'frozen', False) and os.path.exists(msi_path):
                print(f"检测到 Pandoc 已安装，正在清理安装包: {msi_path}")
                os.remove(msi_path)
        except Exception as e:
            # 删除失败不影响使用，仅在控制台打印
            print(f"清理安装包失败 (可能是权限不足或文件被占用): {e}")

        return True

    # 2. 如果系统没找到 Pandoc
    # 即使 settings 说安装了，但系统找不到，也视为未安装（防止误删或环境丢失）

    temp_app = QApplication.instance()
    if not temp_app:
        temp_app = QApplication(sys.argv)

    icon_path = resource_path("icon.ico")
    if os.path.exists(icon_path):
        temp_app.setWindowIcon(QIcon(icon_path))

    msg_box = QMessageBox()
    msg_box.setWindowTitle("组件缺失")
    msg_box.setText("检测到您的电脑尚未安装 Pandoc (文档转换引擎)。\n\n是否立即安装？")
    msg_box.setInformativeText(
        "• 点击【安装】：将自动运行安装程序。安装完成后请**手动重新打开**本软件。\n• 点击【取消】：退出软件。")
    msg_box.setIcon(QMessageBox.Icon.Warning)

    btn_install = msg_box.addButton("安装 Pandoc", QMessageBox.ButtonRole.AcceptRole)
    btn_cancel = msg_box.addButton("取消", QMessageBox.ButtonRole.RejectRole)

    msg_box.exec()

    if msg_box.clickedButton() == btn_install:
        msi_source = resource_path("pandoc.msi")
        temp_dir = tempfile.gettempdir()
        msi_target = os.path.join(temp_dir, "pandoc_install.msi")

        try:
            if os.path.exists(msi_source):
                # 复制到临时目录运行，防止直接运行带来的路径权限问题
                shutil.copy(msi_source, msi_target)
                # 启动安装程序
                subprocess.Popen(["msiexec", "/i", msi_target])

                # [关键修改] 这里不再标记为已安装！
                # 也不删除文件！
                # 直接返回 False 退出软件，等待用户安装完成后重启
                return False
            else:
                # 这种情况通常发生在第二次启动，安装包已经被删除了，但系统里还是没找到 pandoc
                # 提示用户去官网下载
                QMessageBox.critical(None, "安装包未找到",
                                     f"内置安装包已被清理，请前往 pandoc.org 手动下载安装。\n路径: {msi_source}")
                return False
        except Exception as e:
            QMessageBox.critical(None, "错误", f"无法启动安装程序: {str(e)}")
            return False
    else:
        return False


startup_config = load_settings_raw()
if not startup_config.get("enable_gpu", False):
    os.environ["QTWEBENGINE_CHROMIUM_FLAGS"] = "--disable-gpu"
else:
    if "QTWEBENGINE_CHROMIUM_FLAGS" in os.environ:
        del os.environ["QTWEBENGINE_CHROMIUM_FLAGS"]

HAS_DOCX_LIB = False
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm  # [新增] Cm 用于精确控制缩进
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    HAS_DOCX_LIB = True
except ImportError:
    HAS_DOCX_LIB = False


class SettingsDialog(QDialog):
    def __init__(self, parent=None, current_settings=None):
        super().__init__(parent)
        self.setWindowTitle("设置")
        self.setFixedWidth(450)
        self.setWindowFlags(Qt.WindowType.Dialog | Qt.WindowType.WindowCloseButtonHint)
        self.settings = current_settings or {}
        main_layout = QVBoxLayout()
        main_layout.setSpacing(15)

        group_general = QGroupBox("常规设置")
        layout_general = QVBoxLayout()
        lbl_path = QLabel("默认导出位置:")
        path_row = QHBoxLayout()
        self.edit_path = QLineEdit()
        self.edit_path.setPlaceholderText("留空则默认使用上次保存的位置...")
        self.edit_path.setText(self.settings.get("default_export_path", ""))
        btn_browse = QPushButton("📂")
        btn_browse.setFixedWidth(40)
        btn_browse.clicked.connect(self.browse_folder)
        path_row.addWidget(self.edit_path)
        path_row.addWidget(btn_browse)
        layout_general.addWidget(lbl_path)
        layout_general.addLayout(path_row)
        group_general.setLayout(layout_general)
        main_layout.addWidget(group_general)

        group_pdf = QGroupBox("PDF 与 打印设置")
        group_layout = QVBoxLayout()
        row_v = QHBoxLayout()
        lbl_v = QLabel("上下页边距:")
        self.spin_v = QDoubleSpinBox()
        self.spin_v.setRange(0.0, 10.0)
        self.spin_v.setSingleStep(0.1)
        self.spin_v.setSuffix(" cm")
        self.spin_v.setValue(self.settings.get("margin_v", 2.0))
        btn_reset_v = QPushButton("重置")
        btn_reset_v.setFixedWidth(50)
        btn_reset_v.clicked.connect(lambda: self.spin_v.setValue(2.0))
        row_v.addWidget(lbl_v)
        row_v.addWidget(self.spin_v)
        row_v.addWidget(btn_reset_v)
        group_layout.addLayout(row_v)

        row_h = QHBoxLayout()
        lbl_h = QLabel("左右页边距:")
        self.spin_h = QDoubleSpinBox()
        self.spin_h.setRange(0.0, 10.0)
        self.spin_h.setSingleStep(0.1)
        self.spin_h.setSuffix(" cm")
        self.spin_h.setValue(self.settings.get("margin_h", 2.0))
        btn_reset_h = QPushButton("重置")
        btn_reset_h.setFixedWidth(50)
        btn_reset_h.clicked.connect(lambda: self.spin_h.setValue(2.0))
        row_h.addWidget(lbl_h)
        row_h.addWidget(self.spin_h)
        row_h.addWidget(btn_reset_h)
        group_layout.addLayout(row_h)

        self.chk_preview = QCheckBox("在预览界面模拟页边距效果")
        self.chk_preview.setChecked(self.settings.get("show_preview_margins", True))
        group_layout.addWidget(self.chk_preview)
        group_pdf.setLayout(group_layout)
        main_layout.addWidget(group_pdf)

        group_adv = QGroupBox("高级设置 (实验性)")
        group_adv.setStyleSheet("QGroupBox { border: 1px solid #d73a49; } QGroupBox::title { color: #d73a49; }")
        layout_adv = QVBoxLayout()
        self.chk_gpu = QCheckBox("启用 GPU 硬件加速")
        self.chk_gpu.setChecked(self.settings.get("enable_gpu", False))
        lbl_warning = QLabel("⚠️ 警告：开启 GPU 加速可能导致预览区域黑屏。\n(更改此选项需要重启软件生效)")
        lbl_warning.setStyleSheet("color: #666; font-size: 12px; margin-left: 20px;")
        layout_adv.addWidget(self.chk_gpu)
        layout_adv.addWidget(lbl_warning)
        group_adv.setLayout(layout_adv)
        main_layout.addWidget(group_adv)

        main_layout.addStretch()
        btn_layout = QHBoxLayout()
        btn_layout.addStretch()
        self.btn_save = QPushButton("保存设置")
        self.btn_save.setStyleSheet("""
            QPushButton { background-color: #0366d6; color: white; border: none; padding: 6px 25px; border-radius: 4px; font-weight: bold; }
            QPushButton:hover { background-color: #0255b3; }
        """)
        self.btn_save.clicked.connect(self.accept)
        btn_layout.addWidget(self.btn_save)

        main_layout.addLayout(btn_layout)
        self.setLayout(main_layout)

    def browse_folder(self):
        folder = QFileDialog.getExistingDirectory(self, "选择默认输出文件夹")
        if folder:
            self.edit_path.setText(folder)

    def get_data(self):
        return {
            "default_export_path": self.edit_path.text().strip(),
            "margin_v": self.spin_v.value(),
            "margin_h": self.spin_h.value(),
            "show_preview_margins": self.chk_preview.isChecked(),
            "enable_gpu": self.chk_gpu.isChecked()
        }


class TableInsertDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("插入表格")
        self.setFixedWidth(300)
        self.setWindowFlags(Qt.WindowType.Dialog | Qt.WindowType.WindowCloseButtonHint)

        layout = QVBoxLayout()
        form_layout = QFormLayout()

        self.spin_rows = QSpinBox()
        self.spin_rows.setRange(1, 100)
        self.spin_rows.setValue(3)

        self.spin_cols = QSpinBox()
        self.spin_cols.setRange(1, 20)
        self.spin_cols.setValue(3)

        form_layout.addRow("行数 (Rows):", self.spin_rows)
        form_layout.addRow("列数 (Cols):", self.spin_cols)

        layout.addLayout(form_layout)

        self.buttons = QDialogButtonBox(QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel)
        self.buttons.accepted.connect(self.accept)
        self.buttons.rejected.connect(self.reject)

        layout.addWidget(self.buttons)
        self.setLayout(layout)

    def get_data(self):
        return self.spin_rows.value(), self.spin_cols.value()


class SmartEditor(QTextEdit):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.cursorPositionChanged.connect(self.check_task_toggle)
        self.is_fixing = False
        self.textChanged.connect(self.auto_fix_spacing)

    def auto_fix_spacing(self):
        if self.is_fixing: return

        doc = self.document()
        block = doc.begin()

        insert_positions = []
        prev_type = 'empty'
        prev_indent = 0
        in_code_block = False

        while block.isValid():
            text = block.text()

            if text.strip().startswith('```'):
                in_code_block = not in_code_block
                prev_type = 'code_fence'
                prev_indent = 0
                block = block.next()
                continue

            if in_code_block:
                block = block.next()
                continue

            curr_type = self.identify_block_type(text)
            curr_indent = self.get_indent(text)

            if self.check_conflict(prev_type, curr_type, prev_indent, curr_indent):
                insert_positions.append(block.position())

            if curr_type != 'empty':
                prev_type = curr_type
                prev_indent = curr_indent
            else:
                prev_type = 'empty'
                prev_indent = 0

            block = block.next()

        if insert_positions:
            self.is_fixing = True

            cursor = self.textCursor()
            old_pos = cursor.position()

            cursor.beginEditBlock()
            for pos in reversed(insert_positions):
                cursor.setPosition(pos)
                cursor.insertText('\n')
                if pos <= old_pos:
                    old_pos += 1
            cursor.endEditBlock()

            cursor.setPosition(old_pos)
            self.setTextCursor(cursor)
            self.is_fixing = False

    def identify_block_type(self, text):
        clean_text = text.lstrip()
        stripped_text = clean_text.strip()
        if not stripped_text: return 'empty'
        if re.match(r'^#{1,6}\s', clean_text): return 'text'
        if re.match(r'^[-*]\s\[[xX ]\]', clean_text): return 'task'
        if re.match(r'^[-*+]\s', clean_text): return 'ul'
        if re.match(r'^\d+\.', clean_text): return 'ol'
        if re.match(r'^[-*_]{3,}$', stripped_text): return 'hr'
        if stripped_text.startswith('|'): return 'table'
        return 'text'

    def get_indent(self, text):
        match = re.match(r'^(\s*)', text)
        return len(match.group(1)) if match else 0

    def check_conflict(self, t1, t2, i1, i2):
        if t1 == 'empty': return False
        if t2 == 'empty': return False

        list_types = ['ul', 'ol', 'task']

        if t1 == 'text' and t2 in list_types:
            return True
        if t1 in list_types and t2 == 'text':
            if i2 <= i1:
                return True
        if t1 in list_types and t2 in list_types:
            if t1 != t2 and i2 <= i1:
                return True
        if t1 != 'table' and t2 == 'table':
            return True
        if t1 == 'table' and t2 != 'table':
            return True
        return False

    def keyPressEvent(self, event):
        if event.key() == Qt.Key.Key_Return or event.key() == Qt.Key.Key_Enter:
            cursor = self.textCursor()
            block = cursor.block()
            text = block.text()
            task_match = re.match(r'^(\s*)([-*]) \[[xX ]\]\s+', text)
            ul_match = re.match(r'^(\s*)([-*+])\s+', text)
            ol_match = re.match(r'^(\s*)(\d+)\.\s+', text)
            if task_match:
                marker_len = len(task_match.group(0).strip())
                if len(text.strip()) == marker_len:
                    self.remove_current_line_marker(cursor)
                    return
            if ul_match and text.strip() == ul_match.group(2):
                self.remove_current_line_marker(cursor)
                return
            if ol_match and text.strip() == f"{ol_match.group(2)}.":
                self.remove_current_line_marker(cursor)
                return
            if task_match:
                indent = task_match.group(1)
                marker = task_match.group(2)
                self.insert_new_line_with_prefix(f"{indent}{marker} [ ] ")
                return
            elif ul_match:
                indent = ul_match.group(1)
                marker = ul_match.group(2)
                self.insert_new_line_with_prefix(f"{indent}{marker} ")
                return
            elif ol_match:
                indent = ol_match.group(1)
                number = int(ol_match.group(2))
                self.insert_new_line_with_prefix(f"{indent}{number + 1}. ")
                return
        super().keyPressEvent(event)

    def check_task_toggle(self):
        cursor = self.textCursor()
        if cursor.hasSelection(): return
        block = cursor.block()
        text = block.text()
        pos_in_block = cursor.positionInBlock()
        match = re.search(r'^(\s*[-*]\s*)\[([ xX]*)\]', text)
        if match:
            bracket_start = len(match.group(1))
            bracket_end = match.end() - 1
            if bracket_start < pos_in_block <= bracket_end:
                self.blockSignals(True)
                cursor.beginEditBlock()
                content_inside = match.group(2)
                if not content_inside.strip():
                    cursor.setPosition(block.position() + bracket_start)
                    cursor.setPosition(block.position() + bracket_end + 1, QTextCursor.MoveMode.KeepAnchor)
                    cursor.insertText("[x]")
                else:
                    cursor.setPosition(block.position() + bracket_start)
                    cursor.setPosition(block.position() + bracket_end + 1, QTextCursor.MoveMode.KeepAnchor)
                    cursor.insertText("[ ]")
                cursor.endEditBlock()
                target_pos = block.position() + match.end()
                if target_pos < block.position() + block.length():
                    cursor.setPosition(target_pos)
                else:
                    cursor.setPosition(block.position() + block.length() - 1)
                self.setTextCursor(cursor)
                self.blockSignals(False)
                self.textChanged.emit()

    def remove_current_line_marker(self, cursor):
        cursor.beginEditBlock()
        cursor.select(QTextCursor.SelectionType.BlockUnderCursor)
        cursor.removeSelectedText()
        cursor.insertBlock()
        cursor.endEditBlock()
        self.setTextCursor(cursor)

    def insert_new_line_with_prefix(self, prefix):
        cursor = self.textCursor()
        cursor.beginEditBlock()
        cursor.insertText("\n")
        cursor.insertText(prefix)
        cursor.endEditBlock()
        self.setTextCursor(cursor)
        self.ensureCursorVisible()


class MarkdownEditor(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("PaperSmith V16.8")
        self.resize(1200, 800)

        icon_path = resource_path("icon.ico")
        if os.path.exists(icon_path):
            self.setWindowIcon(QIcon(icon_path))

        self.settings_file = SETTINGS_FILE
        self.settings = load_settings_raw()

        default_settings = {
            "margin_v": 2.0,
            "margin_h": 2.0,
            "show_preview_margins": True,
            "default_export_path": "",
            "enable_gpu": False,
            "sync_scroll": True
        }
        for k, v in default_settings.items():
            if k not in self.settings:
                self.settings[k] = v

        self.current_file_path = None
        self.is_preview_ready = False

        self.central_widget = QWidget()
        self.setCentralWidget(self.central_widget)

        self.init_ui_style()

        self.main_layout = QVBoxLayout(self.central_widget)
        self.main_layout.setContentsMargins(0, 0, 0, 0)
        self.splitter = QSplitter(Qt.Orientation.Horizontal)
        self.main_layout.addWidget(self.splitter)

        self.left_panel = QWidget()
        self.left_layout = QVBoxLayout(self.left_panel)
        self.left_layout.setContentsMargins(15, 15, 15, 15)
        self.left_btn_layout = QHBoxLayout()
        btn_import = QPushButton("📂 导入")
        btn_import.clicked.connect(self.open_file)
        btn_paste = QPushButton("📋 粘贴")
        btn_paste.clicked.connect(self.paste_from_clipboard)
        btn_clear = QPushButton("🗑️ 清空")
        btn_clear.setObjectName("btn_clear")
        btn_clear.clicked.connect(self.clear_editor)
        self.left_btn_layout.addWidget(btn_import)
        self.left_btn_layout.addWidget(btn_paste)
        self.left_btn_layout.addWidget(btn_clear)

        self.btn_sync = QPushButton("🔗 同步")
        self.btn_sync.setCheckable(True)
        self.btn_sync.setChecked(self.settings.get("sync_scroll", True))
        self.btn_sync.setToolTip("开启/关闭预览区跟随滚动")
        self.btn_sync.setStyleSheet("""
            QPushButton:checked { background-color: #e1e4e8; border: 1px solid #0366d6; color: #0366d6; }
        """)
        self.btn_sync.toggled.connect(self.save_sync_state)

        self.left_btn_layout.addWidget(self.btn_sync)
        self.left_btn_layout.addStretch()
        self.left_layout.addLayout(self.left_btn_layout)

        self.toolbar_layout = QHBoxLayout()
        self.toolbar_layout.setSpacing(5)
        self.btn_heading = QPushButton("H")
        self.btn_heading.setObjectName("tool_btn")
        self.btn_heading.setToolTip("设置标题")
        self.btn_heading.clicked.connect(self.show_heading_menu)
        self.btn_bold = QPushButton("B")
        self.btn_bold.setObjectName("tool_btn")
        self.btn_bold.setToolTip("粗体")
        self.btn_bold.clicked.connect(self.set_bold)
        self.btn_italic = QPushButton("I")
        self.btn_italic.setObjectName("tool_btn")
        self.btn_italic.setToolTip("斜体")
        self.btn_italic.setStyleSheet("font-style: italic; font-family: 'Times New Roman'; font-weight: bold;")
        self.btn_italic.clicked.connect(self.set_italic)
        self.btn_strike = QPushButton("S")
        self.btn_strike.setObjectName("btn_strike")
        self.btn_strike.setToolTip("删除线")
        self.btn_strike.clicked.connect(self.set_strikethrough)
        line1 = QFrame()
        line1.setFrameShape(QFrame.Shape.VLine)
        line1.setFrameShadow(QFrame.Shadow.Sunken)
        line1.setStyleSheet("background-color: #d1d5da;")
        self.btn_ul = QPushButton("•")
        self.btn_ul.setObjectName("tool_btn")
        self.btn_ul.setToolTip("无序列表")
        self.btn_ul.clicked.connect(self.set_unordered_list)
        self.btn_ol = QPushButton("1.")
        self.btn_ol.setObjectName("tool_btn")
        self.btn_ol.setToolTip("有序列表")
        self.btn_ol.clicked.connect(self.set_ordered_list)
        self.btn_task = QPushButton("☑")
        self.btn_task.setObjectName("tool_btn")
        self.btn_task.setToolTip("任务列表")
        self.btn_task.clicked.connect(self.set_task_list)

        line2 = QFrame()
        line2.setFrameShape(QFrame.Shape.VLine)
        line2.setFrameShadow(QFrame.Shadow.Sunken)
        line2.setStyleSheet("background-color: #d1d5da;")

        self.btn_hr = QPushButton("—")
        self.btn_hr.setObjectName("tool_btn")
        self.btn_hr.setToolTip("插入分割线")
        self.btn_hr.clicked.connect(self.insert_hr)

        self.btn_table = QPushButton("▦")
        self.btn_table.setObjectName("btn_table")
        self.btn_table.setToolTip("插入表格")
        self.btn_table.clicked.connect(self.insert_table)

        self.btn_image = QPushButton("🖼")
        self.btn_image.setObjectName("tool_btn")
        self.btn_image.setToolTip("插入图片")
        self.btn_image.clicked.connect(self.insert_image)

        self.toolbar_layout.addWidget(self.btn_heading)
        self.toolbar_layout.addWidget(self.btn_bold)
        self.toolbar_layout.addWidget(self.btn_italic)
        self.toolbar_layout.addWidget(self.btn_strike)
        self.toolbar_layout.addWidget(line1)
        self.toolbar_layout.addWidget(self.btn_ul)
        self.toolbar_layout.addWidget(self.btn_ol)
        self.toolbar_layout.addWidget(self.btn_task)
        self.toolbar_layout.addWidget(line2)
        self.toolbar_layout.addWidget(self.btn_hr)
        self.toolbar_layout.addWidget(self.btn_table)
        self.toolbar_layout.addWidget(self.btn_image)

        self.toolbar_layout.addStretch()
        self.left_layout.addLayout(self.toolbar_layout)

        self.editor = SmartEditor()
        self.editor.setFont(QFont("Consolas", 11))
        self.editor.setPlaceholderText("在此输入 Markdown 内容...")

        self.editor.verticalScrollBar().valueChanged.connect(self.sync_by_scrollbar)

        self.left_layout.addWidget(self.editor)
        self.splitter.addWidget(self.left_panel)

        self.right_panel = QWidget()
        self.right_layout = QVBoxLayout(self.right_panel)
        self.right_layout.setContentsMargins(15, 15, 15, 15)
        self.right_btn_layout = QHBoxLayout()
        self.right_btn_layout.setSpacing(10)
        self.btn_settings = QPushButton("⚙")
        self.btn_settings.setObjectName("btn_settings")
        self.btn_settings.setToolTip("设置")
        self.btn_settings.setFont(QFont("Segoe UI Emoji", 16))
        self.btn_settings.clicked.connect(self.open_settings_dialog)
        self.btn_save_md = QPushButton("💾 保存")
        self.btn_save_md.clicked.connect(self.save_file)
        self.btn_copy = QPushButton("📋 复制")
        self.btn_copy.clicked.connect(self.copy_content)

        self.right_btn_layout.addStretch()
        self.right_btn_layout.addWidget(self.btn_copy)
        self.right_btn_layout.addWidget(self.btn_save_md)
        self.right_btn_layout.addWidget(self.btn_settings)

        self.right_layout.addLayout(self.right_btn_layout)
        self.preview = QWebEngineView()

        self.debug_page = DebugWebPage(self.preview)
        self.preview.setPage(self.debug_page)
        self.preview.settings().setAttribute(QWebEngineSettings.WebAttribute.LocalContentCanAccessRemoteUrls, True)
        self.preview.settings().setAttribute(QWebEngineSettings.WebAttribute.LocalContentCanAccessFileUrls, True)

        self.preview.page().setBackgroundColor(Qt.GlobalColor.white)
        self.preview.loadFinished.connect(self.on_preview_loaded)

        self.right_layout.addWidget(self.preview)
        self.splitter.addWidget(self.right_panel)
        self.splitter.setSizes([600, 600])

        self.debounce_timer = QTimer()
        self.debounce_timer.setSingleShot(True)
        self.debounce_timer.setInterval(300)
        self.debounce_timer.timeout.connect(self.update_preview)
        self.editor.textChanged.connect(self.debounce_timer.start)
        self.shortcut_save_as = QShortcut(QKeySequence("F12"), self)
        self.shortcut_save_as.activated.connect(self.save_as_file)

    def save_sync_state(self, checked):
        self.settings["sync_scroll"] = checked
        self.save_settings()

    def on_preview_loaded(self, ok):
        if ok:
            self.is_preview_ready = True
            self.sync_by_scrollbar()
            js = """
            try {
                if (typeof katex !== 'undefined') {
                    document.querySelectorAll('.katex-raw').forEach(function(el) {
                        var tex = el.innerText;
                        var isDisplay = el.getAttribute('data-display') === 'true';
                        try {
                            katex.render(tex, el, {
                                displayMode: isDisplay,
                                throwOnError: false
                            });
                        } catch(err) {
                            console.error('KaTeX render error:', err);
                        }
                    });
                    console.log("[KaTeX] Manual rendering finished.");
                } else {
                    console.error("[KaTeX] Library not loaded (yet)!");
                }
            } catch(e) {
                console.error("JS Error:", e);
            }
            """
            self.preview.page().runJavaScript(js)

    def sync_by_scrollbar(self):
        if not self.btn_sync.isChecked(): return

        vbar = self.editor.verticalScrollBar()
        if vbar.maximum() <= 0: return
        ratio = vbar.value() / vbar.maximum()
        js_code = f"window.scrollTo(0, (document.body.scrollHeight - window.innerHeight) * {ratio});"
        self.preview.page().runJavaScript(js_code)

    def load_external_file(self, file_path):
        if file_path and os.path.exists(file_path):
            self.current_file_path = file_path
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    self.editor.setPlainText(f.read())
                self.statusBar().showMessage(f"已加载: {os.path.basename(file_path)}")
                self.update_preview()
            except Exception as e:
                QMessageBox.critical(self, "错误", f"无法打开文件: {str(e)}")

    def save_settings(self):
        save_settings_raw(self.settings)

    def open_settings_dialog(self):
        dialog = SettingsDialog(self, self.settings)
        if dialog.exec():
            new_settings = dialog.get_data()
            self.settings.update(new_settings)
            self.save_settings()
            self.update_preview()
            self.statusBar().showMessage("设置已保存", 3000)

    def get_css_style(self):
        margin_v = self.settings.get("margin_v", 2.0)
        margin_h = self.settings.get("margin_h", 2.0)
        show_preview = self.settings.get("show_preview_margins", True)
        padding_v = f"{margin_v}cm" if show_preview else "20px"
        padding_h = f"{margin_h}cm" if show_preview else "20px"

        katex_scripts = """
        <link rel="stylesheet" href="https://cdn.staticfile.org/KaTeX/0.16.9/katex.min.css">
        <script src="https://cdn.staticfile.org/KaTeX/0.16.9/katex.min.js"></script>
        """

        return f"""
        {katex_scripts}
        <style>
            @page {{ margin: {margin_v}cm {margin_h}cm; }} 
            body {{ font-family: 'Segoe UI', 'Microsoft YaHei', sans-serif; line-height: 1.6; color: #24292e; background-color: #ffffff; }}
            @media screen {{ body {{ padding: {padding_v} {padding_h}; }} }}
            @media print {{ body {{ margin: 0; padding: 0; }} @page {{ margin: 0; }} }}
            ul, ol {{ margin-top: 2px; margin-bottom: 2px; padding-left: 24px; }}
            li p {{ margin: 0; }}
            ol {{ list-style-type: decimal !important; }}
            ul {{ list-style-type: disc !important; }}
            table {{ border-collapse: collapse; width: 100%; margin: 15px 0; }}
            th, td {{ border: 1px solid #dfe2e5; padding: 8px 15px; text-align: left; }}

            th {{ background-color: #f6f8fa; font-weight: bold; }}

            code {{ background-color: #f6f8fa; padding: 0.2em 0.4em; border-radius: 3px; font-family: 'Consolas', monospace; }}
            input[type=checkbox] {{ margin-right: 8px; vertical-align: middle; transform: scale(1.1); }}
            li:has(input[type="checkbox"]) {{ list-style-type: none !important; margin-left: -1.3em; padding-left: 0; }}
            del {{ text-decoration: line-through; color: #666; }}
            img {{ max-width: 100%; height: auto; display: block; margin: 10px 0; border-radius: 4px; }}
            hr {{ border: 0; border-top: 2px solid #dfe2e5; margin: 20px 0; }}

            .katex-display {{ 
                text-align: center; 
                margin: 1em 0; 
                overflow-x: auto; 
                overflow-y: hidden;
            }}
            div.katex-raw[data-display="true"] {{
                text-align: center;
                display: block;
                margin: 10px 0;
            }}
        </style>
        """

    def init_ui_style(self):
        self.setStyleSheet("""
            QMainWindow, QWidget { background-color: #f6f8fa; color: #24292e; font-family: 'Segoe UI', sans-serif; }
            QPushButton { background-color: #ffffff; border: 1px solid #d1d5da; border-radius: 6px; padding: 5px 12px; color: #24292e; font-weight: 500; }
            QPushButton:hover { background-color: #f3f4f6; border-color: #1b1f2326; }
            QPushButton:pressed { background-color: #ebecf0; }

            QPushButton#tool_btn, QPushButton#btn_strike, QPushButton#btn_table {
                background-color: transparent; 
                border: 1px solid transparent; 
                border-radius: 4px;
                min-width: 30px; max-width: 30px; min-height: 30px; max-height: 30px;
                padding: 0px; margin: 0px;
            }

            QPushButton#tool_btn:hover, QPushButton#btn_strike:hover, QPushButton#btn_table:hover { 
                background-color: #e1e4e8; border: 1px solid #d1d5da; 
            }

            QPushButton#tool_btn { font-weight: bold; font-size: 16px; }
            QPushButton#btn_strike { font-weight: bold; font-size: 16px; text-decoration: line-through; }
            QPushButton#btn_table { font-weight: bold; font-size: 25px; color: #24292e; }

            QPushButton#btn_clear { color: #cb2431; border: 1px solid #d1d5da; }
            QPushButton#btn_clear:hover { background-color: #ffeef0; border-color: #cb2431; }

            QPushButton#btn_settings { background-color: transparent; border: 1px solid transparent; border-radius: 18px; padding: 0; min-width: 36px; max-width: 36px; min-height: 36px; max-height: 36px; font-weight: bold; color: #555; }
            QPushButton#btn_settings:hover { background-color: #e1e4e8; color: #24292e; }
            QTextEdit { background-color: #ffffff; color: #24292e; border: 1px solid #e1e4e8; border-radius: 6px; selection-background-color: #0366d6; selection-color: #ffffff; padding: 10px; }
            QSplitter::handle { background-color: #e1e4e8; }
            QMenu { background-color: #ffffff; border: 1px solid #d1d5da; padding: 4px; }
            QMenu::item { background-color: transparent; color: #24292e; padding: 6px 20px; }
            QMenu::item:selected { background-color: #0366d6; color: #ffffff; }
            QGroupBox { font-weight: bold; border: 1px solid #d1d5da; border-radius: 6px; margin-top: 10px; padding-top: 15px; }
            QGroupBox::title { subcontrol-origin: margin; subcontrol-position: top left; left: 10px; padding: 0 5px; background-color: #f6f8fa; }
        """)

    def apply_list_prefix(self, prefix_type):
        cursor = self.editor.textCursor()
        cursor.beginEditBlock()

        start = cursor.selectionStart()
        cursor.setPosition(start)
        block = cursor.block()
        prev_block = block.previous()

        list_pattern = r'^(\s*)([-*+]|\d+\.|[-*+] \[[ xX]\])\s'

        if prev_block.isValid():
            prev_text = prev_block.text()
            if prev_text.strip() and not re.match(list_pattern, prev_text):
                cursor.movePosition(QTextCursor.MoveOperation.StartOfBlock)
                cursor.insertText('\n')

        start = cursor.selectionStart()
        end = cursor.selectionEnd()
        cursor.setPosition(start)
        cursor.movePosition(QTextCursor.MoveOperation.StartOfBlock)
        while cursor.position() <= end:
            cursor.movePosition(QTextCursor.MoveOperation.EndOfBlock, QTextCursor.MoveMode.KeepAnchor)
            text = cursor.selectedText()
            clean_text = re.sub(r'^(\s*)(\d+\.|[-*+]|[-*+] \[[ xX]\])\s+', r'\1', text)
            if prefix_type == 'ul':
                new_text = "- " + clean_text.lstrip()
            elif prefix_type == 'task':
                new_text = "- [ ] " + clean_text.lstrip()
            elif prefix_type == 'ol':
                new_text = "1. " + clean_text.lstrip()
            else:
                new_text = text
            cursor.insertText(new_text)
            if not cursor.block().next().isValid(): break
            cursor.movePosition(QTextCursor.MoveOperation.NextBlock)
            if cursor.position() > end and not cursor.hasSelection(): break
        cursor.endEditBlock()
        self.editor.setFocus()

    def set_unordered_list(self):
        self.apply_list_prefix('ul')

    def set_ordered_list(self):
        self.apply_list_prefix('ol')

    def set_task_list(self):
        self.apply_list_prefix('task')

    def insert_hr(self):
        cursor = self.editor.textCursor()
        cursor.beginEditBlock()
        cursor.movePosition(QTextCursor.MoveOperation.EndOfLine)

        text_block = cursor.block().text()
        if text_block.strip():
            cursor.insertText("\n\n---\n\n")
        else:
            cursor.insertText("\n---\n\n")

        cursor.endEditBlock()
        self.editor.setTextCursor(cursor)
        self.editor.setFocus()

    def insert_table(self):
        dialog = TableInsertDialog(self)
        if dialog.exec():
            rows, cols = dialog.get_data()

            cursor = self.editor.textCursor()
            cursor.beginEditBlock()

            header_row = "| " + " | ".join([f"标题{i + 1}" for i in range(cols)]) + " |"
            divider_row = "| " + " | ".join(["---"] * cols) + " |"

            table_text = f"\n{header_row}\n{divider_row}\n"

            for _ in range(max(0, rows - 1)):
                row_text = "| " + " | ".join(["   "] * cols) + " |\n"
                table_text += row_text

            cursor.insertText(table_text)
            cursor.endEditBlock()
            self.editor.setTextCursor(cursor)
            self.editor.setFocus()

    def insert_image(self):
        file_path, _ = QFileDialog.getOpenFileName(self, "选择图片", "",
                                                   "Images (*.png *.jpg *.jpeg *.bmp *.gif *.svg)")
        if file_path:
            clean_path = file_path.replace('\\', '/')
            cursor = self.editor.textCursor()
            cursor.insertText(f"![图片描述]({clean_path})")
            self.editor.setFocus()

    def get_line_type(self, text):
        text = text.strip()
        if not text: return 'empty'
        if re.match(r'^[-*] \[[xX ]\]', text): return 'task'
        if re.match(r'^[-*+] ', text): return 'ul'
        if re.match(r'^\d+\.', text): return 'ol'
        return 'text'

    def is_render_conflict(self, type_a, type_b):
        list_types = ['task', 'ul', 'ol']
        if type_a in list_types and type_b in list_types:
            if type_a != type_b: return True
        return False

    def update_preview(self):
        raw_text = self.editor.toPlainText()
        lines = raw_text.split('\n')
        processed_lines = []
        prev_type = 'empty'
        for line in lines:
            curr_type = self.get_line_type(line)
            if self.is_render_conflict(prev_type, curr_type): processed_lines.append('')
            processed_lines.append(line)
            prev_type = curr_type
        md_text = '\n'.join(processed_lines)
        md_text = re.sub(r'~~(.*?)~~', r'<del>\1</del>', md_text)
        pattern = re.compile(r'^(\s*)([-*]) \[([ xX])\] ', re.MULTILINE)

        def replace_task(match):
            indent = match.group(1)
            marker = match.group(2)
            is_checked = 'checked' if match.group(3).lower() == 'x' else ''
            return f'{indent}{marker} <input type="checkbox" disabled {is_checked}> '

        processed_md = pattern.sub(replace_task, md_text)

        # --- 使用 MathExtension 扩展 ---
        html_content = markdown.markdown(processed_md,
                                         extensions=['extra', 'codehilite', 'nl2br', 'sane_lists', MathExtension()])

        # --- [无闪烁更新] ---
        if self.is_preview_ready:
            js_content = json.dumps(html_content)
            js = f"""
            (function(){{
                var savedScrollY = window.scrollY;
                document.body.innerHTML = {js_content};
                window.scrollTo(0, savedScrollY);

                if (typeof katex !== 'undefined') {{
                    document.querySelectorAll('.katex-raw').forEach(function(el) {{
                        var tex = el.innerText;
                        var isDisplay = el.getAttribute('data-display') === 'true';
                        try {{
                            katex.render(tex, el, {{
                                displayMode: isDisplay,
                                throwOnError: false
                            }});
                        }} catch(err) {{ console.error(err); }}
                    }});
                }}
            }})();
            """
            self.preview.page().runJavaScript(js)
        else:
            final_html = f"<!DOCTYPE html><html><head><meta charset='utf-8'>{self.get_css_style()}</head><body>{html_content}</body></html>"
            self.preview.setHtml(final_html, QUrl("file:///"))

    def insert_formatting(self, prefix, suffix):
        cursor = self.editor.textCursor()
        if cursor.hasSelection():
            text = cursor.selectedText()
            cursor.insertText(f"{prefix}{text}{suffix}")
        else:
            cursor.insertText(f"{prefix}{suffix}")
            for _ in range(len(suffix)):
                cursor.movePosition(QTextCursor.MoveOperation.Left)
        self.editor.setTextCursor(cursor)
        self.editor.setFocus()

    def set_bold(self):
        self.insert_formatting("**", "**")

    def set_italic(self):
        self.insert_formatting("*", "*")

    def set_strikethrough(self):
        self.insert_formatting("~~", "~~")

    def show_heading_menu(self):
        menu = QMenu(self)
        for i in range(1, 7):
            action = QAction(f"H{i} 标题", self)
            action.triggered.connect(lambda checked, level=i: self.apply_heading(level))
            menu.addAction(action)
        menu.exec(self.btn_heading.mapToGlobal(QPoint(0, self.btn_heading.height() + 5)))

    def apply_heading(self, level):
        cursor = self.editor.textCursor()
        cursor.beginEditBlock()
        cursor.movePosition(QTextCursor.MoveOperation.StartOfBlock)
        cursor.movePosition(QTextCursor.MoveOperation.EndOfBlock, QTextCursor.MoveMode.KeepAnchor)
        text = cursor.selectedText()
        clean_text = re.sub(r'^#+\s*', '', text)
        new_text = ("#" * level) + " " + clean_text
        cursor.insertText(new_text)
        cursor.endEditBlock()
        self.editor.setTextCursor(cursor)
        self.editor.setFocus()

    def get_suggested_filename(self, extension):
        if self.current_file_path:
            base_name = os.path.splitext(os.path.basename(self.current_file_path))[0]
            return f"{base_name}.{extension}"
        content = self.editor.toPlainText()
        for i in range(1, 7):
            pattern = rf'^{"#" * i}\s+(.+)$'
            match = re.search(pattern, content, re.MULTILINE)
            if match:
                title = match.group(1).strip()
                clean_title = re.sub(r'[\\/:*?"<>|]', '_', title).strip(" .")
                clean_title = clean_title[:50]
                if clean_title: return f"{clean_title}.{extension}"
        timestamp = datetime.now().strftime("%Y%m%d_%H%M")
        return f"Export_{timestamp}.{extension}"

    def clear_editor(self):
        if self.editor.toPlainText().strip():
            reply = QMessageBox.question(self, '确认清空', "确定要清空所有内容吗？\n这将清除当前的文件关联。",
                                         QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
                                         QMessageBox.StandardButton.No)
            if reply == QMessageBox.StandardButton.No: return
        self.editor.clear()
        self.current_file_path = None
        self.statusBar().showMessage("内容已清空")

    def open_file(self):
        file_path, _ = QFileDialog.getOpenFileName(self, "选择文件", "", "Markdown (*.md);;All Files (*)")
        if file_path:
            self.current_file_path = file_path
            with open(file_path, 'r', encoding='utf-8') as f: self.editor.setPlainText(f.read())
            self.statusBar().showMessage(f"已加载: {os.path.basename(file_path)}")

    def paste_from_clipboard(self):
        self.current_file_path = None
        self.editor.insertPlainText(QApplication.clipboard().text())
        self.statusBar().showMessage("内容已粘贴")

    def save_file(self):
        if self.current_file_path:
            try:
                with open(self.current_file_path, 'w', encoding='utf-8') as f:
                    f.write(self.editor.toPlainText())
                self.statusBar().showMessage(f"已保存: {os.path.basename(self.current_file_path)}", 3000)
            except Exception as e:
                QMessageBox.critical(self, "保存失败", f"无法写入文件: {str(e)}")
        else:
            self.save_as_file()

    def save_as_file(self):
        default_name = self.get_suggested_filename("md")
        default_dir = self.settings.get("default_export_path", "")
        if self.current_file_path:
            initial_path = os.path.join(os.path.dirname(self.current_file_path), default_name)
        elif default_dir and os.path.isdir(default_dir):
            initial_path = os.path.join(default_dir, default_name)
        else:
            initial_path = default_name

        filters = "Markdown Files (*.md);;PDF Files (*.pdf);;Word Files (*.docx)"
        file_path, selected_filter = QFileDialog.getSaveFileName(self, "另存为", initial_path, filters)

        if not file_path: return

        if selected_filter == "PDF Files (*.pdf)" or file_path.lower().endswith(".pdf"):
            self.export_pdf(target_path=file_path)
        elif selected_filter == "Word Files (*.docx)" or file_path.lower().endswith(".docx"):
            self.export_word(target_path=file_path)
        else:
            try:
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write(self.editor.toPlainText())
                self.current_file_path = file_path
                self.statusBar().showMessage(f"已另存为: {os.path.basename(file_path)}", 3000)
            except Exception as e:
                QMessageBox.critical(self, "保存失败", f"无法写入文件: {str(e)}")

    def copy_content(self):
        def on_html_ready(html):
            def on_text_ready(text):
                mime_data = QMimeData()
                mime_data.setHtml(html)
                mime_data.setText(text)
                QApplication.clipboard().setMimeData(mime_data)
                self.statusBar().showMessage("预览内容已复制到剪贴板", 2000)

            self.preview.page().toPlainText(on_text_ready)

        self.preview.page().toHtml(on_html_ready)

    def export_pdf_wrapper(self):
        self.export_pdf()

    def export_word_wrapper(self):
        self.export_word()

    def export_pdf(self, target_path=None):
        if not self.editor.toPlainText().strip(): QMessageBox.warning(self, "警告", "内容为空，无法导出！"); return
        if not target_path:
            default_name = self.get_suggested_filename("pdf")
            default_dir = self.settings.get("default_export_path", "")
            initial_path = os.path.join(default_dir, default_name) if default_dir and os.path.isdir(
                default_dir) else default_name
            target_path, _ = QFileDialog.getSaveFileName(self, "导出 PDF", initial_path, "PDF Files (*.pdf)")
        if target_path:
            m_v = self.settings.get("margin_v", 2.0)
            m_h = self.settings.get("margin_h", 2.0)
            margins = QMarginsF(m_h * 10, m_v * 10, m_h * 10, m_v * 10)
            layout = QPageLayout(QPageSize(QPageSize.PageSizeId.A4), QPageLayout.Orientation.Portrait, margins)
            layout.setUnits(QPageLayout.Unit.Millimeter)
            self.preview.page().printToPdf(target_path, layout)
            self.statusBar().showMessage(f"PDF 已生成: {target_path}", 5000)

    def export_word(self, target_path=None):
        if not self.editor.toPlainText().strip(): QMessageBox.warning(self, "警告", "内容为空，无法导出！"); return
        if not HAS_DOCX_LIB: QMessageBox.critical(self, "错误", "无法加载 python-docx 库。"); return
        if not target_path:
            default_name = self.get_suggested_filename("docx")
            default_dir = self.settings.get("default_export_path", "")
            initial_path = os.path.join(default_dir, default_name) if default_dir and os.path.isdir(
                default_dir) else default_name
            target_path, _ = QFileDialog.getSaveFileName(self, "导出 Word", initial_path, "Word Files (*.docx)")
        if not target_path: return

        # [关键修复] 使用 gfm 格式，这对任务列表支持更好
        temp_md = "temp_export_cache.md"
        try:
            with open(temp_md, "w", encoding="utf-8") as f:
                f.write(self.editor.toPlainText())
            pypandoc.convert_file(temp_md, 'docx', outputfile=target_path, format='gfm')

            # --- Word 后处理 ---
            self.post_process_word(target_path)
            self.statusBar().showMessage(f"Word 已生成: {target_path}", 5000)
        except Exception as e:
            QMessageBox.critical(self, "导出错误",
                                 f"导出 Word 失败。\n请确认 Pandoc 已正确安装并重启软件。\n错误信息: {str(e)}")
        finally:
            if os.path.exists(temp_md): os.remove(temp_md)

    def post_process_word(self, docx_path):
        doc = Document(docx_path)

        # --- 1. 删除纯空段落 (解决多余空行) ---
        paragraphs_to_delete = []
        for para in doc.paragraphs:
            # 如果段落文本为空且没有包含图片等内联对象，则标记为删除
            if not para.text.strip() and len(para.runs) == 0:
                paragraphs_to_delete.append(para)

        for para in paragraphs_to_delete:
            p = para._element
            p.getparent().remove(p)
            p._p = p._element = None

        # --- 2. 列表与任务列表样式修复 ---
        for para in doc.paragraphs:
            text = para.text.strip()
            style_name = para.style.name

            is_list_style = 'List' in style_name
            # 识别 Unicode 选框符号 (Pandoc GFM 转换后的特征)
            is_task_like = (len(text) > 0 and (
                        text.startswith('☐') or text.startswith('☒') or text.startswith('[ ]') or text.startswith(
                    '[x]')))

            if is_list_style or is_task_like:
                # [核心修复] 强制清除段后间距
                para.paragraph_format.space_after = Pt(0)
                para.paragraph_format.space_before = Pt(0)

                # [核心修复] 统一列表缩进 (修复箭头间距过大)
                # Word 标准列表缩进: Left=0.74cm, FirstLine=-0.74cm (悬挂)
                if is_list_style:
                    if para.paragraph_format.left_indent is None:
                        para.paragraph_format.left_indent = Cm(0.74)
                        para.paragraph_format.first_line_indent = Cm(-0.74)
                    else:
                        # 保持层级，仅修正悬挂
                        para.paragraph_format.first_line_indent = Cm(-0.74)

                # [核心修复] 任务列表去重影 (去除 Word 自动加的圆点)
                # 如果是任务列表且被识别为列表样式，重置为正文样式但保留缩进
                if is_task_like and 'List' in style_name:
                    para.style = doc.styles['Normal']
                    para.paragraph_format.left_indent = Cm(0.74)
                    para.paragraph_format.first_line_indent = Cm(-0.74)

        # --- 3. 表格样式 (保持原样) ---
        header_bg_color = "F6F8FA"
        for table in doc.tables:
            self.set_table_borders(table)
            for i, row in enumerate(table.rows):
                for cell in row.cells:
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    if i == 0:
                        self.set_cell_shading(cell, header_bg_color)
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs: run.font.bold = True

        doc.save(docx_path)

    def set_table_borders(self, table):
        tbl = table._tbl
        tblPr = tbl.tblPr
        tblBorders = tblPr.find(qn('w:tblBorders'))
        if tblBorders is None: tblBorders = OxmlElement('w:tblBorders'); tblPr.append(tblBorders)
        borders = {'top': 'single', 'left': 'single', 'bottom': 'single', 'right': 'single', 'insideH': 'single',
                   'insideV': 'single'}
        for border_name, border_val in borders.items():
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), border_val)
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), 'auto')
            existing = tblBorders.find(qn(f'w:{border_name}'))
            if existing is not None: tblBorders.remove(existing)
            tblBorders.append(border)

    def set_cell_shading(self, cell, fill_color):
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), fill_color)
        tc_pr.append(shd)


if __name__ == "__main__":
    if sys.platform == 'win32':
        myappid = 'mycompany.markdown_editor.v16.7.installer'  # 版本号更新
        ctypes.windll.shell32.SetCurrentProcessExplicitAppUserModelID(myappid)

    # 检测Pandoc
    if check_pandoc_dependency():
        app = QApplication.instance()
        if not app:
            app = QApplication(sys.argv)

        window = MarkdownEditor()

        # 启动时检查命令行参数
        if len(sys.argv) > 1:
            initial_file = sys.argv[1]
            window.load_external_file(initial_file)

        window.show()
        sys.exit(app.exec())
    else:
        sys.exit(0)